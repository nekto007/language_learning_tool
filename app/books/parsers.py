# app/books/parser.py

import logging
import os
import re
import xml.etree.ElementTree as ET

from bs4 import BeautifulSoup
from werkzeug.utils import secure_filename

logger = logging.getLogger(__name__)


def clean_text(text):
    """Очищает текст от некорректных символов и служебной информации"""
    if not text:
        return ''

    # Удаляем некорректные символы кодировки
    text = text.encode('utf-8', 'ignore').decode('utf-8')

    # Удаляем служебную информацию сайтов
    text = re.sub(r'\s*-\s*[a-zA-Z0-9\-_]+\.(com|ru|org|net)\s*$', '', text)

    # Удаляем лишние пробелы, но сохраняем переносы строк
    text = re.sub(r'[ \t]+', ' ', text).strip()

    return text


def parse_book_file(file_path, file_ext, format_type='enhanced'):
    """
    Парсит файл книги и возвращает HTML-контент с сохранением форматирования

    Args:
        file_path: Путь к файлу книги
        file_ext: Расширение файла (.txt, .fb2, .epub, и т.д.)
        format_type: Тип форматирования (auto, simple, enhanced)

    Returns:
        tuple: (html_content, word_count, unique_words_count)
    """
    file_ext = file_ext.lower()

    if file_ext == '.txt':
        return parse_txt(file_path, format_type)
    elif file_ext == '.fb2':
        return parse_fb2(file_path, format_type)
    elif file_ext == '.epub':
        return parse_epub(file_path, format_type)
    elif file_ext == '.pdf':
        return parse_pdf(file_path, format_type)
    elif file_ext == '.docx':
        return parse_docx(file_path, format_type)
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")


def parse_txt(file_path, format_type):
    """Парсит TXT-файл"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
    except UnicodeDecodeError:
        # Попробуем другие кодировки
        encodings = ['latin-1', 'cp1251', 'windows-1251']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                break
            except UnicodeDecodeError:
                continue

    # Процесс нормализации текста
    # Сначала восстанавливаем правильные переносы абзацев
    content = re.sub(r'\n\s*\n', '\n\n', content)
    
    # Удаляем лишние пробелы между словами, НО сохраняем переносы абзацев
    content = re.sub(r'[ \t]+', ' ', content)  # Заменяем только пробелы и табы

    # Подсчет слов после нормализации
    words = re.findall(r'\b[a-zA-Z]+\b', content.lower())
    word_count = len(words)
    unique_words = len(set(words))

    # Форматирование в зависимости от выбранного типа
    if format_type == 'simple':
        # Простое форматирование с абзацами
        paragraphs = content.split('\n\n')
        normalized_paragraphs = []

        for paragraph in paragraphs:
            # Нормализуем текст абзаца, сохраняя переносы строк внутри абзаца
            normalized_paragraph = re.sub(r'[ \t]+', ' ', paragraph).strip()
            if normalized_paragraph:
                normalized_paragraphs.append(normalized_paragraph)

        html_content = '<p>' + '</p><p>'.join(normalized_paragraphs) + '</p>'

    elif format_type == 'enhanced':
        # Улучшенное форматирование для изучения языка
        paragraphs = content.split('\n\n')
        html_parts = []
        chapter_pattern = re.compile(r'^(?:Chapter|CHAPTER)\s+(\d+|[IVXLCDM]+)(?:\s*[:.-]\s*)?(.*)$', re.IGNORECASE)

        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            # Нормализуем текст абзаца, сохраняя переносы строк внутри абзаца
            normalized_paragraph = re.sub(r'[ \t]+', ' ', paragraph).strip()

            # Проверяем, является ли абзац заголовком главы
            chapter_match = chapter_pattern.match(normalized_paragraph)
            if chapter_match:
                chapter_num = chapter_match.group(1)
                chapter_title = chapter_match.group(2).strip() if chapter_match.group(2) else ""
                if chapter_title:
                    html_parts.append(f'<h2>Chapter {chapter_num}: {chapter_title}</h2>')
                else:
                    html_parts.append(f'<h2>Chapter {chapter_num}</h2>')
            else:
                # Важно: добавляем абзац БЕЗ атрибута class
                html_parts.append(f'<p>{normalized_paragraph}</p>')

        html_content = ''.join(html_parts)

    else:  # auto
        # Автоматическое обнаружение структуры с нормализованными пробелами
        normalized_content = re.sub(r'[ \t]+', ' ', content)  # Сохраняем переносы строк
        html_content = f'<div>{normalized_content}</div>'

    return html_content, word_count, unique_words


def extract_fb2_cover(root, namespace):
    """Извлекает обложку из FB2 файла"""
    try:
        import base64
        
        # Ищем coverpage
        coverpage = root.find('.//' + namespace + 'coverpage')
        if coverpage is None:
            return None
            
        # Ищем изображение обложки
        image_ref = coverpage.find('.//' + namespace + 'image')
        if image_ref is None:
            return None
            
        # Получаем ссылку на изображение
        href = image_ref.get('{http://www.w3.org/1999/xlink}href', '')
        if not href and 'href' in image_ref.attrib:
            href = image_ref.get('href')
            
        if not href:
            return None
            
        # Убираем # из начала href
        if href.startswith('#'):
            href = href[1:]
            
        # Ищем binary элемент с этим id
        binary = root.find(f'.//{namespace}binary[@id="{href}"]')
        if binary is None:
            # Пробуем без namespace
            binary = root.find(f'.//binary[@id="{href}"]')
            
        if binary is None:
            return None
            
        # Декодируем base64
        try:
            image_data = base64.b64decode(binary.text)
            content_type = binary.get('content-type', 'image/jpeg')
            
            return {
                'data': image_data,
                'content_type': content_type,
                'size': len(image_data)
            }
        except Exception as decode_err:
            logger.error(f"Error decoding cover image: {str(decode_err)}")
            return None
            
    except Exception as e:
        logger.error(f"Error extracting FB2 cover: {str(e)}")
        return None


def extract_fb2_metadata(file_path):
    """Извлекает метаданные из FB2 файла"""
    logger.info(f"[FB2_PARSER] Starting FB2 metadata extraction from: {file_path}")
    try:
        # Пробуем разные кодировки
        encodings = ['utf-8', 'windows-1251', 'cp1251', 'latin-1']
        content = None

        for encoding in encodings:
            logger.debug(f"[FB2_PARSER] Trying encoding: {encoding}")
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                logger.info(f"[FB2_PARSER] Successfully read file with encoding: {encoding}")
                break
            except UnicodeDecodeError:
                logger.debug(f"[FB2_PARSER] Failed to read with encoding: {encoding}")
                continue

        if content is None:
            logger.error("[FB2_PARSER] Could not decode FB2 file with any encoding")
            return {'title': '', 'author': ''}

        # Парсим XML из строки
        logger.info("[FB2_PARSER] Parsing XML content")
        root = ET.fromstring(content)

        # Определяем namespace
        namespace = ''
        if root.tag.startswith('{'):
            namespace = root.tag.split('}')[0] + '}'
        logger.info(f"[FB2_PARSER] Using XML namespace: '{namespace}'")

        metadata = {'title': '', 'author': '', 'cover_image': None}

        # Извлекаем информацию из description/title-info
        title_info = root.find('.//' + namespace + 'title-info')
        if title_info is not None:
            # Название книги
            book_title = title_info.find(namespace + 'book-title')
            if book_title is not None and book_title.text:
                metadata['title'] = clean_text(book_title.text)

            # Автор книги - собираем все части имени
            author = title_info.find(namespace + 'author')
            if author is not None:
                author_parts = []

                # Первое имя
                first_name = author.find(namespace + 'first-name')
                if first_name is not None and first_name.text:
                    author_parts.append(clean_text(first_name.text))

                # Отчество
                middle_name = author.find(namespace + 'middle-name')
                if middle_name is not None and middle_name.text:
                    author_parts.append(clean_text(middle_name.text))

                # Фамилия
                last_name = author.find(namespace + 'last-name')
                if last_name is not None and last_name.text:
                    author_parts.append(clean_text(last_name.text))

                if author_parts:
                    metadata['author'] = ' '.join(filter(None, author_parts))

        # Извлекаем обложку (только информация о наличии для metadata API)
        cover_data = extract_fb2_cover(root, namespace)
        if cover_data:
            # Для JSON-сериализации сохраняем только метаинформацию
            metadata['cover_image'] = {
                'content_type': cover_data['content_type'],
                'size': cover_data['size'],
                'has_cover': True
            }

        # Логируем для отладки
        logger.info(f"[FB2_PARSER] Successfully extracted FB2 metadata: title='{metadata['title']}', author='{metadata['author']}', cover={'found' if cover_data else 'not found'}")

        return metadata

    except ET.ParseError as e:
        logger.error(f"[FB2_PARSER] XML parsing error in FB2 file: {str(e)}")
        return {'title': '', 'author': ''}
    except Exception as e:
        logger.error(f"[FB2_PARSER] Error extracting FB2 metadata: {str(e)}")
        logger.error(f"[FB2_PARSER] Exception type: {type(e).__name__}")
        return {'title': '', 'author': ''}


def extract_epub_metadata(file_path):
    """Извлекает метаданные из EPUB файла"""
    try:
        import ebooklib
        from ebooklib import epub

        book = epub.read_epub(file_path)
        metadata = {'title': '', 'author': ''}

        # Извлекаем заголовок
        title = book.get_metadata('DC', 'title')
        if title:
            metadata['title'] = title[0][0]

        # Извлекаем автора
        author = book.get_metadata('DC', 'creator')
        if author:
            metadata['author'] = author[0][0]

        return metadata
    except Exception as e:
        logger.error(f"Error extracting EPUB metadata: {str(e)}")
        return {'title': '', 'author': ''}


def extract_docx_metadata(file_path):
    """Извлекает метаданные из DOCX файла"""
    try:
        from docx import Document

        doc = Document(file_path)
        metadata = {'title': '', 'author': ''}

        # Извлекаем метаданные из свойств документа
        if doc.core_properties.title:
            metadata['title'] = doc.core_properties.title
        if doc.core_properties.author:
            metadata['author'] = doc.core_properties.author

        return metadata
    except Exception as e:
        logger.error(f"Error extracting DOCX metadata: {str(e)}")
        return {'title': '', 'author': ''}


def extract_file_metadata(file_path, file_ext):
    """Универсальная функция для извлечения метаданных из файла"""
    logger.info(f"[METADATA_PARSER] Starting metadata extraction for {file_ext} file: {file_path}")
    file_ext = file_ext.lower()

    if file_ext == '.fb2':
        logger.info("[METADATA_PARSER] Using FB2 metadata extractor")
        return extract_fb2_metadata(file_path)
    elif file_ext == '.epub':
        logger.info("[METADATA_PARSER] Using EPUB metadata extractor")
        return extract_epub_metadata(file_path)
    elif file_ext == '.docx':
        logger.info("[METADATA_PARSER] Using DOCX metadata extractor")
        return extract_docx_metadata(file_path)
    else:
        # Для TXT и других форматов пытаемся извлечь из имени файла
        logger.info(f"[METADATA_PARSER] Unsupported format {file_ext}, extracting title from filename")
        filename = os.path.splitext(os.path.basename(file_path))[0]
        result = {'title': filename, 'author': ''}
        logger.info(f"[METADATA_PARSER] Extracted metadata from filename - Title: '{result['title']}'")
        return result


def parse_fb2(file_path, format_type):
    """Парсит FB2-файл (FictionBook)"""
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()

        namespace = ''
        if root.tag.startswith('{'):
            namespace = root.tag.split('}')[0] + '}'
        # Извлекаем текст из body
        body = root.find('.//' + namespace + 'body')
        if body is None:
            raise ValueError("Could not find body element in FB2 file")

        html_parts = []
        word_count = 0
        unique_words = set()
        # Перебираем секции и получаем текст
        for section in body.findall('.//' + namespace + 'section'):
            # Проверяем, есть ли заголовок
            title = section.find('.//' + namespace + 'title')
            if title is not None:
                title_text = ''.join(title.itertext()).strip()
                title_text = re.sub(r'[ \t]+', ' ', title_text)  # Нормализуем пробелы
                if title_text:
                    html_parts.append(f"<h2>{title_text}</h2>")

            # Обрабатываем абзацы
            for p in section.findall('.//' + namespace + 'p'):
                p_text = ''.join(p.itertext()).strip()
                p_text = re.sub(r'[ \t]+', ' ', p_text)  # Нормализуем пробелы
                if p_text:
                    html_parts.append(f'<p class="book-paragraph">{p_text}</p>')

                    # Подсчет слов
                    words = re.findall(r'\b[a-zA-Z]+\b', p_text.lower())
                    word_count += len(words)
                    unique_words.update(words)

        # Формируем HTML контент
        html_content = ''.join(html_parts)
        return html_content, word_count, len(unique_words)

    except Exception as e:
        logger.error(f"Error parsing FB2 file: {str(e)}")
        # В случае ошибки возвращаем простой текст
        return parse_txt(file_path, format_type)


def parse_epub(file_path, format_type):
    """
    Парсит EPUB-файл

    Требует установки библиотеки ebooklib:
    pip install ebooklib
    """
    try:
        import ebooklib
        from ebooklib import epub

        book = epub.read_epub(file_path)

        html_parts = []
        word_count = 0
        unique_words = set()

        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                content = item.get_content().decode('utf-8')

                # Используем BeautifulSoup для парсинга HTML
                soup = BeautifulSoup(content, 'html.parser')

                # Извлекаем заголовки и абзацы
                for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
                    text = tag.get_text().strip()
                    text = re.sub(r'[ \t]+', ' ', text)  # Нормализуем пробелы
                    if text:
                        if tag.name.startswith('h'):
                            html_parts.append(f"<{tag.name}>{text}</{tag.name}>")
                        else:
                            html_parts.append(f'<p class="book-paragraph">{text}</p>')

                        # Подсчет слов
                        words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
                        word_count += len(words)
                        unique_words.update(words)

        # Формируем HTML контент
        html_content = ''.join(html_parts)

        return html_content, word_count, len(unique_words)

    except ImportError:
        logger.error("ebooklib not installed. Install with: pip install ebooklib")
        # В случае отсутствия библиотеки, возвращаем простой текст
        return parse_txt(file_path, format_type)
    except Exception as e:
        logger.error(f"Error parsing EPUB file: {str(e)}")
        return parse_txt(file_path, format_type)


def parse_pdf(file_path, format_type):
    """
    Парсит PDF-файл

    Требует установки библиотеки pypdf:
    pip install pypdf
    """
    try:
        from pypdf import PdfReader

        with open(file_path, 'rb') as file:
            reader = PdfReader(file)

            # Извлекаем текст из всех страниц
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n\n"

        # Нормализация пробелов
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Сначала восстанавливаем абзацы
        text = re.sub(r'[ \t]+', ' ', text)  # Затем нормализуем пробелы

        # Подсчет слов
        words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
        word_count = len(words)
        unique_words = len(set(words))

        # Форматируем текст
        if format_type == 'simple':
            paragraphs = text.split('\n\n')
            normalized_paragraphs = []

            for paragraph in paragraphs:
                normalized_paragraph = re.sub(r'[ \t]+', ' ', paragraph).strip()
                if normalized_paragraph:
                    normalized_paragraphs.append(normalized_paragraph)

            html_content = '<p>' + '</p><p>'.join(normalized_paragraphs) + '</p>'

        elif format_type == 'enhanced':
            paragraphs = text.split('\n\n')
            html_parts = []

            for paragraph in paragraphs:
                paragraph = paragraph.strip()
                if not paragraph:
                    continue

                # Нормализуем текст
                normalized_paragraph = re.sub(r'[ \t]+', ' ', paragraph).strip()

                # Проверяем, может ли абзац быть заголовком
                if len(paragraph) < 100 and paragraph.isupper():
                    html_parts.append(f"<h2>{normalized_paragraph}</h2>")
                else:
                    html_parts.append(f'<p class="book-paragraph">{normalized_paragraph}</p>')

            html_content = ''.join(html_parts)
        else:  # auto
            html_content = f'<div class="book-text">{text}</div>'

        return html_content, word_count, unique_words

    except ImportError:
        logger.error("pypdf not installed. Install with: pip install pypdf")
        return parse_txt(file_path, format_type)
    except Exception as e:
        logger.error(f"Error parsing PDF file: {str(e)}")
        return parse_txt(file_path, format_type)


def parse_docx(file_path, format_type):
    """
    Парсит DOCX-файл

    Требует установки библиотеки python-docx:
    pip install python-docx
    """
    try:
        from docx import Document

        doc = Document(file_path)

        html_parts = []
        word_count = 0
        unique_words = set()

        # Обрабатываем абзацы
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Нормализуем текст
            text = re.sub(r'[ \t]+', ' ', text).strip()

            # Проверяем, является ли абзац заголовком
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.replace('Heading', '')
                try:
                    level = int(level)
                    if 1 <= level <= 6:
                        html_parts.append(f"<h{level}>{text}</h{level}>")
                    else:
                        html_parts.append(f"<h2>{text}</h2>")
                except ValueError:
                    html_parts.append(f"<h2>{text}</h2>")
            else:
                html_parts.append(f'<p class="book-paragraph">{text}</p>')

            # Подсчет слов
            words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
            word_count += len(words)
            unique_words.update(words)

        # Формируем HTML контент
        html_content = ''.join(html_parts)

        return html_content, word_count, len(unique_words)

    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return parse_txt(file_path, format_type)
    except Exception as e:
        logger.error(f"Error parsing DOCX file: {str(e)}")
        return parse_txt(file_path, format_type)


def process_uploaded_book(file, title, format_type='enhanced'):
    """
    Обрабатывает загруженный файл книги

    Args:
        file: Объект загруженного файла
        title: Название книги
        format_type: Тип форматирования

    Returns:
        dict: {'content': html_content, 'word_count': word_count, 'unique_words': unique_words}
    """
    try:
        # Проверяем, что file - это объект файла, а не строка
        if not file or not hasattr(file, 'filename'):
            raise ValueError("Invalid file object - missing filename attribute")

        filename = secure_filename(file.filename)
        file_ext = os.path.splitext(filename)[1].lower()

        # Создаем временный файл
        temp_dir = 'app/temp'
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, filename)

        file.save(temp_path)

        # Обрабатываем файл
        html_content, word_count, unique_words = parse_book_file(temp_path, file_ext, format_type)

        # Удаляем временный файл
        if os.path.exists(temp_path):
            os.remove(temp_path)
        return {
            'content': html_content,
            'word_count': word_count,
            'unique_words': unique_words
        }

    except Exception as e:
        logger.error(f"Error processing book file: {str(e)}")
        # Очищаем временные файлы
        if 'temp_path' in locals() and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except:
                pass
        raise e
